"""文档解析工具 - 支持 PDF/DOCX/HTML/TXT/Markdown"""
import re
import os
from typing import Optional


def parse_file(file_path: str, file_type: str) -> str:
    """根据文件类型解析文件内容，返回纯文本"""
    parsers = {
        "pdf": _parse_pdf,
        "docx": _parse_docx,
        "html": _parse_html,
        "htm": _parse_html,
        "txt": _parse_txt,
        "md": _parse_md,
        "markdown": _parse_md,
    }
    parser = parsers.get(file_type.lower())
    if not parser:
        raise ValueError(f"不支持的文件类型: {file_type}")
    return parser(file_path)


def _parse_pdf(file_path: str) -> str:
    import pdfplumber
    text_parts = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n\n".join(text_parts)


def _parse_docx(file_path: str) -> str:
    from docx import Document
    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # 也提取表格内容
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
    return "\n\n".join(paragraphs)


def _parse_html(file_path: str) -> str:
    from bs4 import BeautifulSoup
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        soup = BeautifulSoup(f.read(), "html.parser")
    # 移除 script 和 style 标签
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()
    text = soup.get_text(separator="\n", strip=True)
    return text


def _parse_txt(file_path: str) -> str:
    encodings = ["utf-8", "gbk", "gb2312", "latin-1"]
    for enc in encodings:
        try:
            with open(file_path, "r", encoding=enc) as f:
                return f.read()
        except (UnicodeDecodeError, UnicodeError):
            continue
    raise ValueError("无法识别文件编码")


def _parse_md(file_path: str) -> str:
    """解析 Markdown：读取原文并去除语法符号，保留纯文本内容"""
    text = _parse_txt(file_path)

    # 代码块：保留内容，去掉围栏
    text = re.sub(r"```[a-zA-Z0-9]*\n?", "", text)
    # 图片：![alt](url) -> alt
    text = re.sub(r"!\[([^\]]*)\]\([^)]*\)", r"\1", text)
    # 链接：[text](url) -> text
    text = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", text)
    # 标题井号
    text = re.sub(r"^#{1,6}\s*", "", text, flags=re.MULTILINE)
    # 加粗 / 斜体 / 行内代码
    text = re.sub(r"(\*\*|__|\*|_|`)([^*_`\n]+)\1", r"\2", text)
    # 行首引用符与列表符
    text = re.sub(r"^>\s?", "", text, flags=re.MULTILINE)
    text = re.sub(r"^\s*[-*+]\s+", "", text, flags=re.MULTILINE)
    # 表格竖线
    text = re.sub(r"\|", " ", text)
    text = re.sub(r"^\s*:?-{3,}:?\s*$", "", text, flags=re.MULTILINE)

    return text


def get_file_type(filename: str) -> str:
    """从文件名获取扩展名"""
    ext = os.path.splitext(filename)[1].lower().lstrip(".")
    return ext
